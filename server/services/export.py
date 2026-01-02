"""Document export service for Word, PDF, and Markdown formats."""

import hashlib
import os
import re
import secrets
from datetime import datetime, timedelta
from io import BytesIO
from typing import Literal, Optional

import aiofiles
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from server.config import settings
from server.models.database import Document, ShareLink

ExportFormat = Literal["word", "pdf", "markdown"]


def _get_red_team_summary_text(red_team_report: dict) -> str:
    """Extract or generate a summary string from the red team report.

    The summary field can be either a string or a dict with statistics.
    """
    if not red_team_report:
        return ""

    raw_summary = red_team_report.get("summary", "")
    if isinstance(raw_summary, dict):
        # Convert summary dict to a readable string
        total_critiques = raw_summary.get("total_critiques", 0)
        total_rounds = raw_summary.get("total_rounds", 0)
        resolution_rate = raw_summary.get("resolution_rate", 0.0)
        consensus = "Reached" if raw_summary.get("consensus_reached", False) else "Not Reached"
        return (
            f"Red Team Report: {total_critiques} critiques across {total_rounds} rounds. "
            f"Resolution rate: {resolution_rate:.1f}%. Consensus: {consensus}."
        )
    return raw_summary if raw_summary else ""


def _markdown_to_reportlab(text: str) -> str:
    """Convert markdown formatting to ReportLab XML tags.

    Handles: **bold**, *italic*, inline code, and escapes special characters.
    Headers and lists are handled separately in _parse_markdown_content.
    """
    # First escape XML special characters (but preserve our markers)
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")

    # Convert **bold** to <b>bold</b>
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)

    # Convert *italic* to <i>italic</i> (but not inside bold)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)

    # Convert `code` to <font name="Courier">code</font>
    text = re.sub(r'`([^`]+)`', r'<font name="Courier">\1</font>', text)

    return text


def _parse_markdown_content(content: str, body_style, heading_style, list_style) -> list:
    """Parse markdown content and return a list of ReportLab flowables.

    Handles:
    - Headers (### Header)
    - Bullet lists (- item)
    - Numbered lists (1. item)
    - Regular paragraphs with bold/italic
    """
    flowables = []
    lines = content.split("\n")
    current_list_items = []
    current_list_type = None  # 'bullet' or 'numbered'
    current_paragraph_lines = []

    def flush_paragraph():
        """Add accumulated paragraph lines as a single paragraph."""
        nonlocal current_paragraph_lines
        if current_paragraph_lines:
            para_text = " ".join(current_paragraph_lines)
            para_text = _markdown_to_reportlab(para_text)
            if para_text.strip():
                flowables.append(Paragraph(para_text, body_style))
            current_paragraph_lines = []

    def flush_list():
        """Add accumulated list items as a ListFlowable."""
        nonlocal current_list_items, current_list_type
        if current_list_items:
            is_numbered = current_list_type == 'numbered'
            list_kwargs = {
                'bulletType': '1' if is_numbered else 'bullet',
                'leftIndent': 20,
            }
            # Only add start parameter for numbered lists
            if is_numbered:
                list_kwargs['start'] = 1
            list_flowable = ListFlowable(
                [ListItem(Paragraph(_markdown_to_reportlab(item), list_style))
                 for item in current_list_items],
                **list_kwargs,
            )
            flowables.append(list_flowable)
            current_list_items = []
            current_list_type = None

    for line in lines:
        stripped = line.strip()

        # Skip empty lines but flush accumulated content
        if not stripped:
            flush_paragraph()
            flush_list()
            continue

        # Check for headers (### Header)
        header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if header_match:
            flush_paragraph()
            flush_list()
            header_text = header_match.group(2)
            # Remove any bold markers from header text
            header_text = re.sub(r'\*\*(.+?)\*\*', r'\1', header_text)
            flowables.append(Spacer(1, 8))
            flowables.append(Paragraph(header_text, heading_style))
            continue

        # Check for bullet list items (- item or * item)
        bullet_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if bullet_match:
            flush_paragraph()
            if current_list_type == 'numbered':
                flush_list()
            current_list_type = 'bullet'
            current_list_items.append(bullet_match.group(1))
            continue

        # Check for numbered list items (1. item)
        numbered_match = re.match(r'^\d+\.\s+(.+)$', stripped)
        if numbered_match:
            flush_paragraph()
            if current_list_type == 'bullet':
                flush_list()
            current_list_type = 'numbered'
            current_list_items.append(numbered_match.group(1))
            continue

        # Regular text - accumulate for paragraph
        flush_list()
        current_paragraph_lines.append(stripped)

    # Flush any remaining content
    flush_paragraph()
    flush_list()

    return flowables


class ExportService:
    """Service for exporting documents to various formats."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.export_dir = settings.export_temp_dir

    async def _ensure_export_dir(self) -> None:
        """Ensure export directory exists."""
        os.makedirs(self.export_dir, exist_ok=True)

    async def export_document(
        self,
        document: Document,
        format: ExportFormat,
    ) -> tuple[bytes, str, str]:
        """
        Export a document to the specified format.

        Args:
            document: Document model to export
            format: Export format (word, pdf, markdown)

        Returns:
            Tuple of (file_bytes, filename, content_type)
        """
        await self._ensure_export_dir()

        if format == "word":
            return await self._export_to_word(document)
        elif format == "pdf":
            return await self._export_to_pdf(document)
        elif format == "markdown":
            return await self._export_to_markdown(document)
        else:
            raise ValueError(f"Unsupported export format: {format}")

    async def _export_to_word(self, document: Document) -> tuple[bytes, str, str]:
        """Export document to Word format (.docx)."""
        doc = DocxDocument()

        # Add title
        title = doc.add_heading(document.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata paragraph
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"Document Type: {document.type} | "
            f"Status: {document.status.capitalize()} | "
            f"Confidence: {document.confidence:.1f}%"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Add sections from content
        if document.content and "sections" in document.content:
            for section in document.content["sections"]:
                # Section title (check both 'title' and 'name' keys for compatibility)
                section_title = section.get("title") or section.get("name") or "Untitled Section"
                doc.add_heading(section_title, level=1)

                # Section content
                content = section.get("content", "")
                # Split by paragraphs
                for para_text in content.split("\n\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

                # Section confidence (if available)
                confidence = section.get("confidence")
                if confidence is not None:
                    conf_para = doc.add_paragraph()
                    conf_run = conf_para.add_run(f"Section Confidence: {confidence:.1f}%")
                    conf_run.font.size = Pt(9)
                    conf_run.font.italic = True
                    conf_run.font.color.rgb = RGBColor(100, 100, 100)

                doc.add_paragraph()  # Spacer between sections

        # Add red team summary if available
        summary_text = _get_red_team_summary_text(document.red_team_report)
        if summary_text:
            doc.add_heading("Red Team Analysis Summary", level=1)
            doc.add_paragraph(summary_text)

        # Add generation metrics if available
        if document.metrics:
            doc.add_heading("Generation Metrics", level=1)
            metrics = document.metrics
            metrics_text = (
                f"Rounds Completed: {metrics.get('roundsCompleted', metrics.get('rounds_completed', 0))}\n"
                f"Total Critiques: {metrics.get('totalCritiques', metrics.get('total_critiques', 0))}\n"
                f"Critical Issues: {metrics.get('criticalCount', metrics.get('critical_count', 0))}\n"
                f"Major Issues: {metrics.get('majorCount', metrics.get('major_count', 0))}\n"
                f"Minor Issues: {metrics.get('minorCount', metrics.get('minor_count', 0))}"
            )
            doc.add_paragraph(metrics_text)

        # Add footer with generation timestamp
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            f"Generated on {document.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = self._sanitize_filename(document.title) + ".docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return buffer.getvalue(), filename, content_type

    async def _export_to_pdf(self, document: Document) -> tuple[bytes, str, str]:
        """Export document to PDF format."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        # Get styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            alignment=1,  # Center
            spaceAfter=12,
        )

        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=10,
            textColor=colors.gray,
            alignment=1,  # Center
            spaceAfter=24,
        )

        section_title_style = ParagraphStyle(
            "SectionTitle",
            parent=styles["Heading2"],
            fontSize=14,
            spaceBefore=18,
            spaceAfter=12,
        )

        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=11,
            spaceAfter=12,
            leading=14,
        )

        confidence_style = ParagraphStyle(
            "Confidence",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.gray,
            fontName="Helvetica-Oblique",
            spaceAfter=12,
        )

        # Style for sub-headers within sections (### headers in markdown)
        sub_heading_style = ParagraphStyle(
            "SubHeading",
            parent=styles["Heading3"],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=8,
            fontName="Helvetica-Bold",
        )

        # Style for list items
        list_item_style = ParagraphStyle(
            "ListItem",
            parent=styles["Normal"],
            fontSize=11,
            spaceAfter=4,
            leading=14,
            leftIndent=10,
        )

        story = []

        # Title
        story.append(Paragraph(document.title, title_style))

        # Metadata
        meta_text = (
            f"Document Type: {document.type} | "
            f"Status: {document.status.capitalize()} | "
            f"Confidence: {document.confidence:.1f}%"
        )
        story.append(Paragraph(meta_text, meta_style))

        # Sections
        if document.content and "sections" in document.content:
            for section in document.content["sections"]:
                # Section title (check both 'title' and 'name' keys for compatibility)
                section_title = section.get("title") or section.get("name") or "Untitled Section"
                story.append(Paragraph(section_title, section_title_style))

                # Section content - parse markdown formatting
                content = section.get("content", "")
                content_flowables = _parse_markdown_content(
                    content, body_style, sub_heading_style, list_item_style
                )
                story.extend(content_flowables)

                # Section confidence
                confidence = section.get("confidence")
                if confidence is not None:
                    story.append(
                        Paragraph(f"Section Confidence: {confidence:.1f}%", confidence_style)
                    )

        # Red team summary
        summary_text = _get_red_team_summary_text(document.red_team_report)
        if summary_text:
            story.append(Spacer(1, 24))
            story.append(Paragraph("Red Team Analysis Summary", section_title_style))
            safe_summary = (
                summary_text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(safe_summary, body_style))

        # Metrics table
        if document.metrics:
            story.append(Spacer(1, 24))
            story.append(Paragraph("Generation Metrics", section_title_style))

            metrics = document.metrics
            data = [
                ["Metric", "Value"],
                ["Rounds Completed", str(metrics.get("roundsCompleted", metrics.get("rounds_completed", 0)))],
                ["Total Critiques", str(metrics.get("totalCritiques", metrics.get("total_critiques", 0)))],
                ["Critical Issues", str(metrics.get("criticalCount", metrics.get("critical_count", 0)))],
                ["Major Issues", str(metrics.get("majorCount", metrics.get("major_count", 0)))],
                ["Minor Issues", str(metrics.get("minorCount", metrics.get("minor_count", 0)))],
            ]

            table = Table(data, colWidths=[2.5 * inch, 1.5 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                        ("TEXTCOLOR", (0, 1), (-1, -1), colors.black),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 1), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )
            story.append(table)

        # Footer
        story.append(Spacer(1, 36))
        footer_style = ParagraphStyle(
            "Footer",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.gray,
            alignment=1,
        )
        story.append(
            Paragraph(
                f"Generated on {document.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')}",
                footer_style,
            )
        )

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        filename = self._sanitize_filename(document.title) + ".pdf"
        content_type = "application/pdf"

        return buffer.getvalue(), filename, content_type

    async def _export_to_markdown(self, document: Document) -> tuple[bytes, str, str]:
        """Export document to Markdown format."""
        lines = []

        # Title
        lines.append(f"# {document.title}")
        lines.append("")

        # Metadata
        lines.append(
            f"**Document Type:** {document.type} | "
            f"**Status:** {document.status.capitalize()} | "
            f"**Confidence:** {document.confidence:.1f}%"
        )
        lines.append("")
        lines.append("---")
        lines.append("")

        # Sections
        if document.content and "sections" in document.content:
            for section in document.content["sections"]:
                # Section title (check both 'title' and 'name' keys for compatibility)
                section_title = section.get("title") or section.get("name") or "Untitled Section"
                lines.append(f"## {section_title}")
                lines.append("")

                # Section content
                content = section.get("content", "")
                lines.append(content)
                lines.append("")

                # Section confidence
                confidence = section.get("confidence")
                if confidence is not None:
                    lines.append(f"*Section Confidence: {confidence:.1f}%*")
                    lines.append("")

        # Red team summary
        summary_text = _get_red_team_summary_text(document.red_team_report)
        if summary_text:
            lines.append("---")
            lines.append("")
            lines.append("## Red Team Analysis Summary")
            lines.append("")
            lines.append(summary_text)
            lines.append("")

        # Metrics
        if document.metrics:
            lines.append("---")
            lines.append("")
            lines.append("## Generation Metrics")
            lines.append("")
            metrics = document.metrics
            lines.append("| Metric | Value |")
            lines.append("|--------|-------|")
            lines.append(f"| Rounds Completed | {metrics.get('roundsCompleted', metrics.get('rounds_completed', 0))} |")
            lines.append(f"| Total Critiques | {metrics.get('totalCritiques', metrics.get('total_critiques', 0))} |")
            lines.append(f"| Critical Issues | {metrics.get('criticalCount', metrics.get('critical_count', 0))} |")
            lines.append(f"| Major Issues | {metrics.get('majorCount', metrics.get('major_count', 0))} |")
            lines.append(f"| Minor Issues | {metrics.get('minorCount', metrics.get('minor_count', 0))} |")
            lines.append("")

        # Footer
        lines.append("---")
        lines.append("")
        lines.append(
            f"*Generated on {document.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')}*"
        )

        content = "\n".join(lines)
        filename = self._sanitize_filename(document.title) + ".md"
        content_type = "text/markdown; charset=utf-8"

        return content.encode("utf-8"), filename, content_type

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize a filename for safe file system use."""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        sanitized = filename
        for char in invalid_chars:
            sanitized = sanitized.replace(char, "_")
        # Limit length
        if len(sanitized) > 200:
            sanitized = sanitized[:200]
        return sanitized.strip()


class ShareLinkService:
    """Service for managing document share links."""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def create_share_link(
        self,
        document_id: str,
        password: Optional[str] = None,
        expires_in_days: Optional[int] = None,
    ) -> ShareLink:
        """
        Create a shareable link for a document.

        Args:
            document_id: ID of document to share
            password: Optional password protection
            expires_in_days: Optional expiration in days (None = never expires)

        Returns:
            Created ShareLink model
        """
        # Generate secure token
        token = secrets.token_urlsafe(32)

        # Hash password if provided
        password_hash = None
        if password:
            password_hash = hashlib.sha256(password.encode()).hexdigest()

        # Calculate expiration
        expires_at = None
        if expires_in_days:
            expires_at = datetime.utcnow() + timedelta(days=expires_in_days)

        share_link = ShareLink(
            document_id=document_id,
            token=token,
            password_hash=password_hash,
            expires_at=expires_at,
        )

        self.db.add(share_link)
        await self.db.flush()
        await self.db.refresh(share_link)

        return share_link

    async def get_share_link_by_token(self, token: str) -> Optional[ShareLink]:
        """Get a share link by its token."""
        stmt = select(ShareLink).where(ShareLink.token == token)
        result = await self.db.execute(stmt)
        return result.scalar_one_or_none()

    async def validate_share_link(
        self,
        token: str,
        password: Optional[str] = None,
    ) -> tuple[bool, Optional[str], Optional[ShareLink]]:
        """
        Validate a share link.

        Args:
            token: Share link token
            password: Password if required

        Returns:
            Tuple of (is_valid, error_message, share_link)
        """
        share_link = await self.get_share_link_by_token(token)

        if not share_link:
            return False, "Share link not found", None

        # Check expiration
        if share_link.expires_at and share_link.expires_at < datetime.utcnow():
            return False, "Share link has expired", None

        # Check password
        if share_link.password_hash:
            if not password:
                return False, "Password required", share_link
            if hashlib.sha256(password.encode()).hexdigest() != share_link.password_hash:
                return False, "Invalid password", share_link

        return True, None, share_link

    async def delete_share_link(self, share_link_id: str) -> bool:
        """Delete a share link."""
        stmt = select(ShareLink).where(ShareLink.id == share_link_id)
        result = await self.db.execute(stmt)
        share_link = result.scalar_one_or_none()

        if not share_link:
            return False

        await self.db.delete(share_link)
        await self.db.flush()
        return True

    async def get_share_links_for_document(self, document_id: str) -> list[ShareLink]:
        """Get all share links for a document."""
        stmt = select(ShareLink).where(ShareLink.document_id == document_id)
        result = await self.db.execute(stmt)
        return list(result.scalars().all())
